from flask import Flask, render_template, request, jsonify, session, redirect, url_for, Response, stream_with_context, flash
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import os
import json
import google.generativeai as genai
from google.generativeai.types import GenerationConfig
from datetime import datetime
import docx
import re
import time
from dotenv import load_dotenv
import fitz

load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-super-secret-key-here-meta-chat-ai-future-v2'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///metachatai_future_v2.db'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

try:
    gemini_api_key = os.getenv('GEMINI_API_KEY')
    if not gemini_api_key:
        raise ValueError("Hata: GEMINI_API_KEY ortam değişkeni bulunamadı.")
    genai.configure(api_key=gemini_api_key)
except Exception as e:
    print(f"Gemini API configuration error: {e}")

db = SQLAlchemy(app)

# --- Veritabanı Modelleri ---
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(120), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    assistants = db.relationship('Assistant', backref='owner', lazy=True, cascade="all, delete-orphan")

class Assistant(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text)
    instructions = db.Column(db.Text, nullable=True, default="Sen, son derece zeki, yardımsever ve uzman bir yapay zeka asistanısın.")
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    documents = db.relationship('Document', backref='assistant', lazy=True, cascade="all, delete-orphan")
    conversations = db.relationship('Conversation', backref='assistant', lazy=True, cascade="all, delete-orphan")

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)
    filepath = db.Column(db.String(500), nullable=False)
    filetype = db.Column(db.String(50), nullable=False)
    assistant_id = db.Column(db.Integer, db.ForeignKey('assistant.id'), nullable=False)
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)
    content = db.Column(db.Text)
    file_size = db.Column(db.Integer, default=0)

class Conversation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200))
    assistant_id = db.Column(db.Integer, db.ForeignKey('assistant.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    messages = db.relationship('Message', backref='conversation', lazy=True, cascade="all, delete-orphan")

class Message(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    content = db.Column(db.Text, nullable=False)
    is_user = db.Column(db.Boolean, default=True)
    conversation_id = db.Column(db.Integer, db.ForeignKey('conversation.id'), nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)
    documents = db.Column(db.Text)


# --- Yardımcı Fonksiyonlar ---
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ['pdf', 'txt', 'docx']

def extract_text_from_pdf(filepath):
    try:
        with fitz.open(filepath) as doc:
            return "".join(page.get_text() for page in doc)
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"

def extract_text_from_docx(filepath):
    try:
        doc = docx.Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error extracting DOCX: {str(e)}"

def extract_text_from_txt(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        return f"Error extracting TXT: {str(e)}"

def clean_text(text):
    if not text: return ""
    return re.sub(r'\s+', ' ', text).strip()

def get_gemini_config_and_prompt(prompt, assistant_instructions, context="", chat_history=""):
    model = genai.GenerativeModel('gemini-1.5-pro-latest')
    system_prompt = f"""{assistant_instructions}

    **EK KURALLAR:**
    1. **ÖNCELİK SIRASI:**
        * **Önce BELGELER:** Cevap "YÜKLENEN DOSYA İÇERİKLERİ"nde varsa, cevabını **kesinlikle** bu belgelere dayandır ve hangi belgeden aldığını belirt. (Örn: `rapor.pdf` belgesine göre...).
        * **Sonra GENEL BİLGİ:** Cevap belgelerde yoksa, kendi engin bilgini kullan. Bu durumu mutlaka belirt. (Örn: "Yüklediğiniz belgelerde bu konuyla ilgili bilgi bulamadım, ancak genel bilgilerime göre...").
    2. **FORMATLAMA (ÇOK ÖNEMLİ):** Cevaplarını zenginleştirmek için **Markdown formatını** kullan. Önemli noktaları **kalın** yaz, listeler oluştur, başlıklar ve kod blokları kullan.
    """
    full_prompt = f"{system_prompt}\n\nYÜKLENEN DOSYA İÇERİKLERİ:\n---\n{context if context else 'Yok'}\n---\n\nÖNCEKİ KONUŞMA GEÇMİŞİ:\n---\n{chat_history if chat_history else 'Yok'}\n---\n\nKULLANICI SORUSU: {prompt}\n\nCEVAP:"
    generation_config = GenerationConfig(temperature=0.7)
    return model, full_prompt, generation_config

def get_chat_history(conversation_id, limit=10):
    if not conversation_id: return ""
    messages = Message.query.filter_by(conversation_id=conversation_id).order_by(Message.timestamp.desc()).limit(limit).all()
    return "\n".join([f"{'Kullanıcı' if msg.is_user else 'Asistan'}: {msg.content}" for msg in reversed(messages)])

# --- Ana Route'lar ---

@app.route('/chat/<int:assistant_id>', methods=['GET', 'POST'])
def chat(assistant_id):
    if 'user_id' not in session: return redirect(url_for('login'))
    assistant = Assistant.query.get_or_404(assistant_id)
    if assistant.owner.id != session['user_id']: return redirect(url_for('dashboard'))

    if request.method == 'POST':
        data = request.get_json()
        message_content = data.get('message')
        conversation_id = data.get('conversation_id')
        attached_documents = data.get('attached_documents', [])
        
        new_conversation_id = None
        if not conversation_id:
            title = message_content[:50].strip()
            conversation = Conversation(title=title, assistant_id=assistant_id)
            db.session.add(conversation)
            db.session.commit()
            conversation_id = conversation.id
            new_conversation_id = conversation.id
        
        documents_json = json.dumps(attached_documents) if attached_documents else None
        user_message = Message(content=message_content, is_user=True, conversation_id=conversation_id, documents=documents_json)
        db.session.add(user_message)
        db.session.commit()

        documents = Document.query.filter(Document.assistant_id == assistant_id, Document.id.in_(attached_documents)).all() if attached_documents else Document.query.filter_by(assistant_id=assistant_id).all()
        context = "\n\n".join([f"--- Belge: {doc.filename} ---\n{doc.content}" for doc in documents if doc.content])
        chat_history = get_chat_history(conversation_id)

        def generate_response_chunks():
            full_response_text = ""
            try:
                model, full_prompt, config = get_gemini_config_and_prompt(message_content, assistant.instructions, context, chat_history)
                stream = model.generate_content(full_prompt, generation_config=config, stream=True)
                for chunk in stream:
                    if chunk.text:
                        full_response_text += chunk.text
                        yield chunk.text
            except Exception as e:
                yield f"⚠️ Bir hata oluştu: {str(e)}"
            finally:
                if full_response_text:
                    with app.app_context():
                        ai_message = Message(content=full_response_text.strip(), is_user=False, conversation_id=conversation_id)
                        db.session.add(ai_message)
                        db.session.commit()
        
        response = Response(stream_with_context(generate_response_chunks()), mimetype='text/plain')
        if new_conversation_id:
            response.headers['X-Conversation-Id'] = new_conversation_id
        return response

    conversation_id = request.args.get('conversation_id')
    messages = []
    if conversation_id:
        conv = Conversation.query.get(conversation_id)
        if conv and conv.assistant_id == assistant_id:
            messages = Message.query.filter_by(conversation_id=conversation_id).order_by(Message.timestamp).all()
    
    documents = Document.query.filter_by(assistant_id=assistant_id).order_by(Document.uploaded_at.desc()).all()
    
    conversations = Conversation.query.filter_by(assistant_id=assistant_id).order_by(Conversation.created_at.desc()).all()

    for message in messages:
        if message.documents:
            try: message.documents_for_display = json.loads(message.documents)
            except: message.documents_for_display = []
        else: message.documents_for_display = []
            
    return render_template('chat.html', assistant=assistant, messages=messages, documents=documents, conversation_id=conversation_id, conversations=conversations)


@app.route('/assistant/<int:assistant_id>/edit', methods=['GET', 'POST'])
def edit_assistant(assistant_id):
    if 'user_id' not in session: return redirect(url_for('login'))
    assistant = Assistant.query.get_or_404(assistant_id)
    if assistant.owner.id != session['user_id']:
        flash('Bu asistana erişim yetkiniz yok.', 'danger')
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        assistant.name = request.form.get('name', assistant.name)
        assistant.description = request.form.get('description', assistant.description)
        assistant.instructions = request.form.get('instructions', assistant.instructions)
        db.session.commit()
        flash('Asistan başarıyla güncellendi!', 'success')
        return redirect(url_for('assistant_detail', assistant_id=assistant.id))
    
    return render_template('edit_assistant.html', assistant=assistant)


@app.route('/conversation/<int:conversation_id>/rename', methods=['POST'])
def rename_conversation(conversation_id):
    if 'user_id' not in session: return jsonify({'error': 'Unauthorized'}), 401
    
    conversation = Conversation.query.get_or_404(conversation_id)
    if conversation.assistant.owner.id != session['user_id']:
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.get_json()
    new_title = data.get('title', '').strip()
    
    if new_title and len(new_title) <= 200:
        conversation.title = new_title
        db.session.commit()
        return jsonify({'message': 'Sohbet başarıyla yeniden adlandırıldı.'})
    
    return jsonify({'error': 'Başlık boş olamaz veya çok uzun.'}), 400


@app.route('/conversation/<int:conversation_id>/delete', methods=['DELETE'])
def delete_conversation(conversation_id):
    if 'user_id' not in session: return jsonify({'error': 'Unauthorized'}), 401
    
    conv = Conversation.query.get_or_404(conversation_id)
    if conv.assistant.owner.id != session['user_id']:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        db.session.delete(conv)
        db.session.commit()
        return jsonify({'message': 'Sohbet başarıyla silindi.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        if User.query.filter_by(username=username).first():
            flash('Bu kullanıcı adı zaten alınmış', 'danger')
            return redirect(url_for('register'))
        if User.query.filter_by(email=email).first():
            flash('Bu e-posta zaten kayıtlı', 'danger')
            return redirect(url_for('register'))
        user = User(username=username, email=email, password_hash=generate_password_hash(password))
        db.session.add(user)
        db.session.commit()
        session['user_id'] = user.id
        flash('Hesabınız başarıyla oluşturuldu!', 'success')
        return redirect(url_for('dashboard'))
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password_hash, password):
            session['user_id'] = user.id
            return redirect(url_for('dashboard'))
        else:
            flash('Geçersiz kullanıcı adı veya şifre.', 'danger')
            return redirect(url_for('login'))
    return render_template('login.html')

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session: return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    assistants = Assistant.query.filter_by(user_id=user.id).all()
    return render_template('dashboard.html', user=user, assistants=assistants)

@app.route('/create_assistant', methods=['GET', 'POST'])
def create_assistant():
    if 'user_id' not in session: return redirect(url_for('login'))
    if request.method == 'POST':
        name = request.form.get('name')
        description = request.form.get('description')
        assistant = Assistant(name=name, description=description, user_id=session['user_id'])
        db.session.add(assistant)
        db.session.commit()
        flash('Yeni asistan başarıyla oluşturuldu!', 'success')
        return redirect(url_for('assistant_detail', assistant_id=assistant.id))
    return render_template('create_assistant.html')

@app.route('/assistant/<int:assistant_id>')
def assistant_detail(assistant_id):
    if 'user_id' not in session: return redirect(url_for('login'))
    assistant = Assistant.query.get_or_404(assistant_id)
    if assistant.owner.id != session['user_id']:
        flash('Bu asistana erişim yetkiniz yok.', 'danger')
        return redirect(url_for('dashboard'))
    conversations = Conversation.query.filter_by(assistant_id=assistant_id).order_by(Conversation.created_at.desc()).all()
    documents = Document.query.filter_by(assistant_id=assistant_id).order_by(Document.uploaded_at.desc()).all()
    return render_template('assistant_detail.html', assistant=assistant, conversations=conversations, documents=documents)

# <<< YENİ: Sohbet listesini dinamik olarak döndüren route >>>
@app.route('/assistant/<int:assistant_id>/conversations')
def get_conversations(assistant_id):
    if 'user_id' not in session: return jsonify({'error': 'Unauthorized'}), 401
    assistant = Assistant.query.get_or_404(assistant_id)
    if assistant.owner.id != session['user_id']:
        return jsonify({'error': 'Unauthorized'}), 403
    
    conversations = Conversation.query.filter_by(assistant_id=assistant_id).order_by(Conversation.created_at.desc()).all()
    # Sadece sohbet listesini içeren küçük bir HTML parçasını render et
    return render_template('_conversation_list.html', conversations=conversations, assistant=assistant)

@app.route('/upload_document/<int:assistant_id>', methods=['POST'])
def upload_document(assistant_id):
    if 'user_id' not in session: return jsonify({'error': 'Unauthorized'}), 401
    assistant = Assistant.query.get_or_404(assistant_id)
    if assistant.owner.id != session['user_id']: return jsonify({'error': 'Unauthorized'}), 403
    if 'file' not in request.files: return jsonify({'error': 'Dosya seçilmedi'}), 400
    file = request.files['file']
    if file.filename == '': return jsonify({'error': 'Dosya seçilmedi'}), 400
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], str(assistant_id), filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        file.save(filepath)
        file_extension = filename.rsplit('.', 1)[1].lower()
        content = ""
        if file_extension == 'pdf': content = extract_text_from_pdf(filepath)
        elif file_extension == 'docx': content = extract_text_from_docx(filepath)
        elif file_extension == 'txt': content = extract_text_from_txt(filepath)
        document = Document(filename=filename, filepath=filepath, filetype=file_extension, assistant_id=assistant_id, content=clean_text(content), file_size=os.path.getsize(filepath))
        db.session.add(document)
        db.session.commit()
        return jsonify({'message': 'Dosya başarıyla yüklendi!', 'document_id': document.id, 'filename': filename, 'file_size': document.file_size})
    return jsonify({'error': 'Geçersiz dosya türü.'}), 400

@app.route('/get_documents/<int:assistant_id>')
def get_documents(assistant_id):
    if 'user_id' not in session: return jsonify({'error': 'Unauthorized'}), 401
    assistant = Assistant.query.get_or_404(assistant_id)
    if assistant.owner.id != session['user_id']: return jsonify({'error': 'Unauthorized'}), 403
    docs = Document.query.filter_by(assistant_id=assistant_id).order_by(Document.uploaded_at.desc()).all()
    return jsonify({'documents': [{'id': d.id, 'filename': d.filename, 'filetype': d.filetype, 'uploaded_at': d.uploaded_at.strftime('%d.%m.%Y %H:%M'), 'file_size': d.file_size or 0} for d in docs]})

@app.route('/delete_document/<int:document_id>', methods=['DELETE'])
def delete_document(document_id):
    if 'user_id' not in session: return jsonify({'error': 'Unauthorized'}), 401
    doc = Document.query.get_or_404(document_id)
    if doc.assistant.owner.id != session['user_id']: return jsonify({'error': 'Unauthorized'}), 403
    try:
        if os.path.exists(doc.filepath): os.remove(doc.filepath)
        db.session.delete(doc)
        db.session.commit()
        return jsonify({'message': 'Dosya başarıyla silindi'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/logout')
def logout():
    session.pop('user_id', None)
    flash('Başarıyla çıkış yaptınız.', 'success')
    return redirect(url_for('index'))

@app.route('/reset-db')
def reset_db():
    if 'user_id' not in session: return redirect(url_for('login'))
    db.drop_all()
    db.create_all()
    flash('Veritabanı başarıyla sıfırlandı!', 'warning')
    return redirect(url_for('dashboard'))

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)